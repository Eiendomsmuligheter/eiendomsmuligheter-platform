from typing import Dict, List, Optional, Any, Union
import logging
from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfkit
from datetime import datetime
from dataclasses import dataclass
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image as PILImage
import io
import emoji
from typing import Dict, List, Optional, Any
from weasyprint import HTML, CSS
import os

logger = logging.getLogger(__name__)

@dataclass
class DocumentTemplate:
    """Utvidet dokumentmalklasse med rike formateringsalternativer"""
    name: str
    sections: List[str]
    required_data: List[str]
    format: str
    language: str = "nb_NO"
    style_settings: Dict[str, Any] = None
    header_image: Optional[str] = None
    footer_text: Optional[str] = None
    page_numbers: bool = True
    include_toc: bool = True
    watermark: Optional[str] = None
    fonts: Dict[str, str] = None
    color_scheme: Dict[str, str] = None
    emojis: bool = True

class DocumentGenerator:
    """Avansert dokumentgenerator med rike formateringsalternativer og AI-støtte"""
    
    def __init__(self, config_path: Optional[str] = None):
        """Initialiserer dokumentgenerator med avanserte innstillinger"""
        self.templates = self._load_templates()
        self.current_doc = None
        self.style_manager = self._initialize_style_manager()
        self.chart_generator = self._initialize_chart_generator()
        self.emoji_support = self._initialize_emoji_support()
        self.config = self._load_config(config_path)
        self.language_manager = self._initialize_language_manager()
        
    def _load_templates(self) -> Dict[str, DocumentTemplate]:
        """Last inn utvidede dokumentmaler med rik formatering"""
        return {
            "bruksendring": DocumentTemplate(
                name="Søknad om bruksendring",
                sections=[
                    "forside",
                    "sammendrag",
                    "introduksjon",
                    "eiendomsinfo",
                    "teknisk_beskrivelse",
                    "brannsikring",
                    "ventilasjon",
                    "støy",
                    "parkering",
                    "avfall",
                    "energi",
                    "miljøpåvirkning",
                    "universell_utforming",
                    "kostnader",
                    "tidsplan",
                    "konklusjon",
                    "vedlegg"
                ],
                required_data=[
                    "eiendom_id",
                    "eier_info",
                    "tekniske_spesifikasjoner",
                    "brannkrav",
                    "ventilasjonskrav"
                ],
                format="docx",
                style_settings={
                    "theme": "modern",
                    "color_scheme": "professional",
                    "font_family": "Calibri",
                    "heading_style": "modern"
                },
                header_image="logo.png",
                footer_text="Generert av EiendomsAI Pro - Bygger fremtidens boliger",
                page_numbers=True,
                include_toc=True,
                emojis=True
            ),
            "teknisk_rapport": DocumentTemplate(
                name="Teknisk rapport",
                sections=[
                    "forside",
                    "sammendrag",
                    "bygningsbeskrivelse",
                    "teknisk_tilstand",
                    "utbedringstiltak",
                    "kostnadsestimater",
                    "energianalyse",
                    "miljøvurdering",
                    "vedlegg"
                ],
                required_data=[
                    "bygningsdata",
                    "tilstandsvurdering",
                    "kostnader",
                    "energidata",
                    "miljødata"
                ],
                format="pdf",
                style_settings={
                    "theme": "technical",
                    "color_scheme": "professional",
                    "font_family": "Arial",
                    "heading_style": "technical"
                }
            )
        }

    def generate_document(
        self,
        template_name: str,
        data: Dict[str, Any],
        output_format: Optional[str] = None,
        style_override: Optional[Dict[str, Any]] = None,
        include_charts: bool = True,
        include_3d: bool = True,
        quality: str = "high"
    ) -> Path:
        """
        Genererer et rikt formatert dokument med avanserte funksjoner
        
        Args:
            template_name: Navn på malen som skal brukes
            data: Data som skal inkluderes i dokumentet
            output_format: Ønsket output format (docx/pdf/html)
            style_override: Overstyring av standard stiler
            include_charts: Inkluder datavisualiseringer
            include_3d: Inkluder 3D-visualiseringer
            quality: Kvalitetsnivå for bilder og grafer
        """
        try:
            # Valider og forbered data
            template = self._get_template(template_name)
            self._validate_data(template, data)
            
            # Generer dokument basert på format
            if template.format == "docx" or output_format == "docx":
                doc_path = self._generate_rich_docx(template, data, style_override)
            elif template.format == "pdf" or output_format == "pdf":
                doc_path = self._generate_rich_pdf(template, data, style_override)
            else:
                raise ValueError(f"Ukjent dokumentformat: {template.format}")
            
            # Legg til ekstra funksjoner
            if include_charts:
                self._add_data_visualizations(doc_path, data)
            if include_3d:
                self._add_3d_visualizations(doc_path, data)
            
            # Optimaliser output
            final_path = self._optimize_document(doc_path, quality)
            
            return final_path
            
        except Exception as e:
            logger.error(f"Feil ved dokumentgenerering: {str(e)}")
            raise
            
    def _generate_rich_docx(
        self,
        template: DocumentTemplate,
        data: Dict[str, Any],
        style_override: Optional[Dict[str, Any]] = None
    ) -> Path:
        """Genererer et rikt formatert Word-dokument med avanserte funksjoner"""
        try:
            doc = Document()
            
            # Konfigurer dokumentstiler
            self._configure_document_styles(doc, template, style_override)
            
            # Legg til forside
            self._add_cover_page(doc, template, data)
            
            # Legg til innholdsfortegnelse
            if template.include_toc:
                self._add_table_of_contents(doc)
            
            # Legg til seksjoner med rik formatering
            for section in template.sections:
                self._add_rich_section(doc, section, data, template)
            
            # Legg til sidenummerering og topptekst/bunntekst
            self._add_page_numbers(doc)
            self._add_headers_and_footers(doc, template)
            
            # Lagre dokument
            output_path = self._get_output_path(template)
            doc.save(str(output_path))
            
            return output_path
            
        except Exception as e:
            logger.error(f"Feil ved generering av rikt Word-dokument: {str(e)}")
            raise
            
    def _add_rich_section(
        self,
        doc: Document,
        section: str,
        data: Dict[str, Any],
        template: DocumentTemplate
    ) -> None:
        """Legger til en rikt formatert seksjon i dokumentet"""
        # Legg til seksjonsoverskrift med stil
        heading = doc.add_heading(
            self._format_section_title(section),
            level=1
        )
        self._apply_heading_style(heading, template)
        
        # Legg til innhold basert på seksjonstype
        if section == "sammendrag":
            self._add_executive_summary(doc, data)
        elif section == "teknisk_beskrivelse":
            self._add_technical_description(doc, data)
        elif section == "kostnadsestimater":
            self._add_cost_analysis(doc, data)
        # ... flere seksjonstyper ...
        
        # Legg til relevant visualisering
        if section in self._get_visualization_sections():
            self._add_section_visualization(doc, section, data)
            
    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legger til et profesjonelt sammendrag med hovedpunkter"""
        # Legg til sammendragstekst
        summary = doc.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Legg til hovedpunkter i en pen formatering
        if "hovedpunkter" in data:
            for punkt in data["hovedpunkter"]:
                bullet = summary.add_run(f"• {punkt}\n")
                bullet.font.color.rgb = RGBColor(0, 112, 192)
                
    def _add_cost_analysis(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legger til detaljert kostnadsanalyse med grafer og tabeller"""
        if "kostnader" not in data:
            return
            
        costs = data["kostnader"]
        
        # Legg til kostnadssammendrag
        table = doc.add_table(rows=1, cols=3)
        self._style_table(table, "LightShading-Accent1")
        
        # Fyll tabell med data
        for category, amount in costs.items():
            row = table.add_row().cells
            row[0].text = category
            row[1].text = f"{amount:,.2f} NOK"
            row[2].text = self._calculate_percentage(amount, sum(costs.values()))
            
        # Legg til kostnadsdiagram
        self._add_cost_chart(doc, costs)
        
    def generate_municipality_documents(
        self,
        property_data: Dict[str, Any],
        municipality: str
    ) -> List[Path]:
        """Genererer alle nødvendige dokumenter for kommunal søknad"""
        try:
            # Hent kommunespesifikke krav
            requirements = self._get_municipality_requirements(municipality)
            
            # Generer alle påkrevde dokumenter
            documents = []
            for doc_type in requirements["required_documents"]:
                template = self._get_municipality_template(doc_type, municipality)
                doc_path = self.generate_document(
                    template_name=template,
                    data=self._prepare_municipality_data(property_data, doc_type)
                )
                documents.append(doc_path)
            
            return documents
            
        except Exception as e:
            logger.error(f"Feil ved generering av kommunale dokumenter: {str(e)}")
            raise

    def create_presentation(
        self,
        data: Dict[str, Any],
        template: str = "modern",
        include_3d: bool = True
    ) -> Path:
        """Genererer en profesjonell presentasjon av prosjektet"""
        try:
            # Initialiser presentasjon
            pres = self._initialize_presentation(template)
            
            # Legg til innhold
            self._add_title_slide(pres, data)
            self._add_overview_slide(pres, data)
            self._add_technical_slides(pres, data)
            
            if include_3d:
                self._add_3d_visualization_slides(pres, data)
            
            # Lagre presentasjon
            output_path = self._get_output_path("presentation")
            pres.save(str(output_path))
            
            return output_path
            
        except Exception as e:
            logger.error(f"Feil ved generering av presentasjon: {str(e)}")
            raise

    def _style_table(self, table: Any, style: str) -> None:
        """Anvender avansert tabellformatering"""
        table.style = style
        table.autofit = True
        
        # Legg til avansert formatering
        for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Sett kantlinjer
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                
                for border in ['top', 'left', 'bottom', 'right']:
                    edge = OxmlElement(f'w:{border}')
                    edge.set(qn('w:val'), 'single')
                    edge.set(qn('w:sz'), '4')
                    edge.set(qn('w:space'), '0')
                    edge.set(qn('w:color'), '4F81BD')
                    tcBorders.append(edge)
                
                tcPr.append(tcBorders)
                
    def _calculate_percentage(self, value: float, total: float) -> str:
        """Beregner og formaterer prosentandel"""
        if total == 0:
            return "0%"
        percentage = (value / total) * 100
        return f"{percentage:.1f}%"