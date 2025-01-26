from typing import Dict, List, Optional, Any
import logging
from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
from datetime import datetime
from dataclasses import dataclass

logger = logging.getLogger(__name__)

@dataclass
class DocumentTemplate:
    name: str
    sections: List[str]
    required_data: List[str]
    format: str
    language: str = "nb_NO"

class DocumentGenerator:
    def __init__(self):
        self.templates = self._load_templates()
        self.current_doc = None
        
    def _load_templates(self) -> Dict[str, DocumentTemplate]:
        """Last inn dokumentmaler"""
        return {
            "bruksendring": DocumentTemplate(
                name="Søknad om bruksendring",
                sections=[
                    "introduksjon",
                    "eiendomsinfo",
                    "teknisk_beskrivelse",
                    "brannsikring",
                    "ventilasjon",
                    "støy",
                    "parkering",
                    "avfall",
                    "konklusjon"
                ],
                required_data=[
                    "eiendom_id",
                    "eier_info",
                    "tekniske_spesifikasjoner",
                    "brannkrav",
                    "ventilasjonskrav"
                ],
                format="docx"
            ),
            "teknisk_rapport": DocumentTemplate(
                name="Teknisk rapport",
                sections=[
                    "sammendrag",
                    "bygningsbeskrivelse",
                    "teknisk_tilstand",
                    "utbedringstiltak",
                    "kostnadsestimater"
                ],
                required_data=[
                    "bygningsdata",
                    "tilstandsvurdering",
                    "kostnader"
                ],
                format="pdf"
            )
        }
        
    def generate_document(self, template_name: str, data: Dict[str, Any]) -> Path:
        """Generer dokument basert på mal og data"""
        try:
            template = self.templates.get(template_name)
            if not template:
                raise ValueError(f"Fant ikke mal: {template_name}")
                
            # Valider data
            self._validate_data(template, data)
            
            # Opprett dokument
            if template.format == "docx":
                return self._generate_docx(template, data)
            elif template.format == "pdf":
                return self._generate_pdf(template, data)
            else:
                raise ValueError(f"Ukjent dokumentformat: {template.format}")
                
        except Exception as e:
            logger.error(f"Feil ved dokumentgenerering: {str(e)}")
            return None
            
    def _validate_data(self, template: DocumentTemplate, data: Dict[str, Any]) -> bool:
        """Valider at all nødvendig data er tilgjengelig"""
        missing = [field for field in template.required_data if field not in data]
        if missing:
            raise ValueError(f"Mangler påkrevde felt: {', '.join(missing)}")
        return True
        
    def _generate_docx(self, template: DocumentTemplate, data: Dict[str, Any]) -> Path:
        """Generer Word-dokument"""
        try:
            doc = Document()
            
            # Sett dokumentegenskaper
            doc.core_properties.title = template.name
            doc.core_properties.author = "EiendomsAI Pro"
            doc.core_properties.created = datetime.now()
            
            # Legg til tittel
            title = doc.add_heading(template.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Legg til seksjoner
            for section in template.sections:
                self._add_section(doc, section, data)
                
            # Lagre dokument
            output_path = Path(f"/tmp/{template.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(str(output_path))
            
            return output_path
            
        except Exception as e:
            logger.error(f"Feil ved generering av Word-dokument: {str(e)}")
            return None
            
    def _generate_pdf(self, template: DocumentTemplate, data: Dict[str, Any]) -> Path:
        """Generer PDF-dokument"""
        try:
            # Først generer som docx
            docx_path = self._generate_docx(template, data)
            if not docx_path:
                raise ValueError("Kunne ikke generere Word-dokument")
                
            # Konverter til PDF
            pdf_path = docx_path.with_suffix('.pdf')
            # Konverter til PDF med pdfkit
            options = {
                'page-size': 'A4',
                'margin-top': '20mm',
                'margin-right': '20mm',
                'margin-bottom': '20mm',
                'margin-left': '20mm',
                'encoding': 'UTF-8',
                'enable-local-file-access': None
            }
            
            pdfkit.from_file(str(docx_path), str(pdf_path), options=options)
            
            # Slett midlertidig docx-fil
            docx_path.unlink()
            
            return pdf_path
            
        except Exception as e:
            logger.error(f"Feil ved generering av PDF: {str(e)}")
            return None
            
    def _add_section(self, doc: Document, section: str, data: Dict[str, Any]) -> None:
        """Legg til en seksjon i dokumentet"""
        # Legg til overskrift
        doc.add_heading(section.replace('_', ' ').title(), level=1)
        
        # Legg til innhold basert på seksjonstype
        if section == "introduksjon":
            self._add_introduction(doc, data)
        elif section == "eiendomsinfo":
            self._add_property_info(doc, data)
        elif section == "teknisk_beskrivelse":
            self._add_technical_description(doc, data)
        elif section == "brannsikring":
            self._add_fire_safety(doc, data)
        elif section == "ventilasjon":
            self._add_ventilation(doc, data)
        elif section == "kostnadsestimater":
            self._add_cost_estimates(doc, data)
        elif section == "støy":
            self._add_noise_requirements(doc, data)
        elif section == "parkering":
            self._add_parking_requirements(doc, data)
        elif section == "avfall":
            self._add_waste_management(doc, data)
        elif section == "konklusjon":
            self._add_conclusion(doc, data)
        
        # Legg til mellomrom
        doc.add_paragraph()
        
    def _add_introduction(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til introduksjonsseksjon"""
        p = doc.add_paragraph()
        p.add_run("Dette dokumentet er generert av EiendomsAI Pro. ")
        p.add_run("Dato: ").bold = True
        p.add_run(datetime.now().strftime("%d.%m.%Y"))
        
    def _add_property_info(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til eiendomsinformasjon"""
        if "eiendom_id" in data:
            p = doc.add_paragraph()
            p.add_run("Eiendomsinformasjon").bold = True
            p.add_run(f"\nEiendoms-ID: {data['eiendom_id']}")
            
    def _add_technical_description(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til teknisk beskrivelse"""
        if "tekniske_spesifikasjoner" in data:
            p = doc.add_paragraph()
            p.add_run("Teknisk Beskrivelse").bold = True
            specs = data["tekniske_spesifikasjoner"]
            
            # Areal og romfordeling
            if "areal" in specs:
                p = doc.add_paragraph()
                p.add_run("Areal og romfordeling:").bold = True
                p.add_run(f"\nBruttoareal: {specs['areal'].get('brutto', 'N/A')} m²")
                p.add_run(f"\nNettoareal: {specs['areal'].get('netto', 'N/A')} m²")
                p.add_run(f"\nAntall rom: {specs['areal'].get('antall_rom', 'N/A')}")
            
            # Konstruksjon og materialer
            if "konstruksjon" in specs:
                p = doc.add_paragraph()
                p.add_run("Konstruksjon og materialer:").bold = True
                for key, value in specs["konstruksjon"].items():
                    p.add_run(f"\n- {key}: {value}")
            
            # Tekniske installasjoner
            if "installasjoner" in specs:
                p = doc.add_paragraph()
                p.add_run("Tekniske installasjoner:").bold = True
                for inst in specs["installasjoner"]:
                    p.add_run(f"\n- {inst}")
                    
    def _add_fire_safety(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til brannsikring informasjon"""
        if "brannkrav" in data:
            p = doc.add_paragraph()
            p.add_run("Brannsikring").bold = True
            fire_reqs = data["brannkrav"]
            
            for req in fire_reqs:
                p = doc.add_paragraph(req["beskrivelse"], style='List Bullet')
                if "tiltak" in req:
                    for tiltak in req["tiltak"]:
                        doc.add_paragraph(tiltak, style='List Bullet 2')
                        
    def _add_ventilation(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til ventilasjonsinformasjon"""
        if "ventilasjonskrav" in data:
            p = doc.add_paragraph()
            p.add_run("Ventilasjon").bold = True
            vent_reqs = data["ventilasjonskrav"]
            
            for rom, krav in vent_reqs.items():
                p = doc.add_paragraph()
                p.add_run(f"Rom: {rom}").bold = True
                p.add_run(f"\nLuftveksling: {krav.get('luftveksling', 'N/A')}")
                p.add_run(f"\nKrav til avtrekk: {krav.get('avtrekk', 'N/A')}")
                
    def _add_cost_estimates(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til kostnadsestimater"""
        if "kostnader" in data:
            p = doc.add_paragraph()
            p.add_run("Kostnadsestimater").bold = True
            costs = data["kostnader"]
            
            total = 0
            for kategori, beløp in costs.items():
                p = doc.add_paragraph()
                p.add_run(f"{kategori}: ").bold = True
                p.add_run(f"{beløp:,.2f} NOK")
                total += beløp
            
            p = doc.add_paragraph()
            p.add_run("Totalt estimert kostnad: ").bold = True
            p.add_run(f"{total:,.2f} NOK")
            
    def _add_noise_requirements(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til støykrav og tiltak"""
        if "støykrav" in data:
            p = doc.add_paragraph()
            p.add_run("Støykrav og tiltak").bold = True
            noise_reqs = data["støykrav"]
            
            if "grenseverdier" in noise_reqs:
                p = doc.add_paragraph()
                p.add_run("Grenseverdier:").bold = True
                for type_støy, grense in noise_reqs["grenseverdier"].items():
                    p.add_run(f"\n- {type_støy}: {grense} dB")
            
            if "tiltak" in noise_reqs:
                p = doc.add_paragraph()
                p.add_run("Nødvendige tiltak:").bold = True
                for tiltak in noise_reqs["tiltak"]:
                    p.add_run(f"\n- {tiltak}")
                    
    def _add_parking_requirements(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til parkeringskrav"""
        if "parkering" in data:
            p = doc.add_paragraph()
            p.add_run("Parkeringskrav").bold = True
            parking = data["parkering"]
            
            if "krav" in parking:
                p = doc.add_paragraph()
                p.add_run(f"Minimumskrav: {parking['krav']} plasser")
            
            if "tilgjengelig" in parking:
                p = doc.add_paragraph()
                p.add_run(f"Tilgjengelige plasser: {parking['tilgjengelig']}")
                
            if "tiltak" in parking:
                p = doc.add_paragraph()
                p.add_run("Nødvendige tiltak:").bold = True
                for tiltak in parking["tiltak"]:
                    p.add_run(f"\n- {tiltak}")
                    
    def _add_waste_management(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til avfallshåndtering"""
        if "avfall" in data:
            p = doc.add_paragraph()
            p.add_run("Avfallshåndtering").bold = True
            waste = data["avfall"]
            
            if "krav" in waste:
                p = doc.add_paragraph()
                p.add_run("Kommunale krav:").bold = True
                for krav in waste["krav"]:
                    p.add_run(f"\n- {krav}")
            
            if "løsning" in waste:
                p = doc.add_paragraph()
                p.add_run("Planlagt løsning:").bold = True
                p.add_run(f"\n{waste['løsning']}")
                
    def _add_conclusion(self, doc: Document, data: Dict[str, Any]) -> None:
        """Legg til konklusjon"""
        p = doc.add_paragraph()
        p.add_run("Konklusjon").bold = True
        
        if "konklusjon" in data:
            konk = data["konklusjon"]
            
            if "hovedpunkter" in konk:
                p = doc.add_paragraph()
                p.add_run("Hovedpunkter:").bold = True
                for punkt in konk["hovedpunkter"]:
                    p.add_run(f"\n- {punkt}")
            
            if "anbefalinger" in konk:
                p = doc.add_paragraph()
                p.add_run("\nAnbefalinger:").bold = True
                for anbefaling in konk["anbefalinger"]:
                    p.add_run(f"\n- {anbefaling}")
            
            if "neste_steg" in konk:
                p = doc.add_paragraph()
                p.add_run("\nNeste steg:").bold = True
                for steg in konk["neste_steg"]:
                    p.add_run(f"\n- {steg}")
                    
            # Legg til dato og signatur
            p = doc.add_paragraph()
            p.add_run("\n\nDato: ").bold = True
            p.add_run(datetime.now().strftime("%d.%m.%Y"))
            
            p = doc.add_paragraph()
            p.add_run("\nGenerert av EiendomsAI Pro")